from django.shortcuts import render
from django.http import JsonResponse,FileResponse,HttpResponse
from django.conf import settings
# Create your views here.
from docx import Document
import io
import os
def index(request):
    return render(request, 'index.html')

import json
from django.http import JsonResponse


def send_data(request):
    if request.method == 'POST':
        d1 = request.POST.get('d1')
        print("ffffffffff",d1)
        print(type(d1))
        response_data = {
            'message': 'Data received successfully!',  # You can customize this message
            'd1': d1  # You can include other data if needed
        }
        return HttpResponse(json.dumps(response_data), content_type='application/json')
    else:
        return HttpResponse(status=400)  # Return a 400 Bad Request status for non-POST requests

from io import BytesIO
def download_docx(request):
    tab1_content = request.POST.get('tab1_content', '')
    tab2_content = request.POST.get('tab2_content', '')

    print("tab1content", tab1_content)

    doc = Document()
    doc.add_paragraph("Tab 1 Content:")
    doc.add_paragraph(tab1_content)
    doc.add_paragraph("Tab 2 Content:")
    doc.add_paragraph(tab2_content)
    buffer = BytesIO()
    doc.save(buffer)

    # Set the buffer position to the beginning
    buffer.seek(0)

    # Create a response with the buffer content as the file
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="data.docx"'

    # Close the buffer to release resources
    buffer.close()

    return response